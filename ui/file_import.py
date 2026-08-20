"""
ElderCare - File import helpers.

Tries to extract text from PDF/Word/Excel uploads, then heuristically
matches it against the disease/medication knowledge banks so the user can
confirm and import codes into the elder profile.

This is a best-effort extractor - the user always confirms before saving.
"""

from __future__ import annotations
from pathlib import Path
import io
import re


def extract_text_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def extract_text_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def extract_text_xlsx(file_bytes: bytes) -> str:
    try:
        import pandas as pd
        sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        out = []
        for name, df in sheets.items():
            out.append(f"== {name} ==")
            out.append(df.astype(str).to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        return f"[XLSX extraction failed: {e}]"


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(file_bytes)
    if ext == ".docx":
        return extract_text_docx(file_bytes)
    if ext in (".xlsx", ".xls"):
        return extract_text_xlsx(file_bytes)
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def match_codes_in_text(text: str, conn) -> dict:
    """
    Heuristic match: scan text for disease/medication/condition names
    (both Hebrew and English). Returns candidates by table.
    """
    text_low = text.lower()
    results = {"diseases": [], "medications": [], "conditions": []}

    for table, key in [
        ("kb_diseases", "diseases"),
        ("kb_medications", "medications"),
        ("kb_conditions", "conditions"),
    ]:
        rows = conn.execute(
            f"SELECT code, name_he, name_en FROM {table}"
        ).fetchall()
        for r in rows:
            for name in (r["name_he"], r["name_en"]):
                if not name:
                    continue
                # boundary-ish match; Hebrew has no word boundaries in regex's sense
                if name and name.lower() in text_low:
                    results[key].append({
                        "code": r["code"],
                        "name_he": r["name_he"],
                        "name_en": r["name_en"],
                    })
                    break
    return results
