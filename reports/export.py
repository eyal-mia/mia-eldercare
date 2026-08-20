"""
ElderCare - Report exporters (Word / PDF / Excel).

Each entrypoint takes (conn, elder_id, days_back, lang) and returns the path
to a file written under exports/.
"""

from __future__ import annotations
from pathlib import Path
import datetime as dt
import sys

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import pandas as pd
from ui.i18n import t

EXPORT_DIR = ROOT / "exports"
EXPORT_DIR.mkdir(exist_ok=True)


def _summary_frame(conn, elder_id: int, days_back: int) -> dict:
    elder = conn.execute("SELECT * FROM elders WHERE id = ?", (elder_id,)).fetchone()
    cutoff = (dt.date.today() - dt.timedelta(days=days_back)).isoformat()

    plans = pd.read_sql(
        "SELECT * FROM daily_plans WHERE elder_id = ? AND plan_date >= ? ORDER BY plan_date",
        conn, params=(elder_id, cutoff),
    )
    items = pd.read_sql("""
        SELECT pi.*, dp.plan_date, ka.name_he, ka.name_en
        FROM plan_items pi
        JOIN daily_plans dp ON dp.id = pi.plan_id
        LEFT JOIN kb_activities ka ON ka.code = pi.activity_code
        WHERE dp.elder_id = ? AND dp.plan_date >= ?
        ORDER BY dp.plan_date, pi.start_time
    """, conn, params=(elder_id, cutoff))
    measurements = pd.read_sql(
        "SELECT * FROM measurements WHERE elder_id = ? AND measurement_date >= ? "
        "ORDER BY measurement_date",
        conn, params=(elder_id, cutoff),
    )

    total = len(items)
    done = int(items["executed"].sum()) if total else 0
    rate = (done / total * 100) if total else 0.0
    return {
        "elder": dict(elder),
        "cutoff": cutoff,
        "plans": plans,
        "items": items,
        "measurements": measurements,
        "total_items": total,
        "done_items": done,
        "completion_rate": rate,
    }


def _filename(elder, ext: str) -> Path:
    safe = "".join(c for c in elder["full_name"] if c.isalnum() or c in " -_")
    stamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    return EXPORT_DIR / f"{safe.replace(' ', '_')}-{stamp}.{ext}"


# ---------- WORD ----------
def elder_summary_docx(conn, elder_id: int, days_back: int, lang: str) -> str:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = _summary_frame(conn, elder_id, days_back)
    elder = data["elder"]
    name_field = "name_he" if lang == "he" else "name_en"

    doc = Document()
    title = doc.add_heading(f"{t('report_summary', lang)} - {elder['full_name']}", level=0)
    if lang == "he":
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    info = doc.add_paragraph()
    info.add_run(f"{t('birth_date', lang)}: {elder.get('birth_date') or '-'}    ")
    info.add_run(f"{t('room_number', lang)}: {elder.get('room_number') or '-'}    ")
    info.add_run(f"{t('report_period_days', lang)}: {days_back}")

    doc.add_heading(t("report_executions", lang), level=1)
    p = doc.add_paragraph()
    p.add_run(f"{t('items_scheduled', lang)}: {data['total_items']}\n")
    p.add_run(f"{t('executed_total', lang)}: {data['done_items']}\n")
    p.add_run(f"{t('completion_rate', lang)}: {data['completion_rate']:.1f}%\n")

    if not data["items"].empty:
        doc.add_heading(t("nav_daily_plan", lang), level=1)
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = t("plan_for_date", lang).strip()
        hdr[1].text = t("time_morning", lang).split()[0]
        hdr[2].text = t("nav_profile", lang)
        hdr[3].text = t("duration_min", lang)
        hdr[4].text = t("executed", lang)
        for _, row in data["items"].iterrows():
            r = table.add_row().cells
            r[0].text = str(row["plan_date"])
            r[1].text = row["time_slot"] or ""
            r[2].text = str(row[name_field] or row["activity_code"])
            r[3].text = str(row["duration_min"] or "")
            r[4].text = t("yes", lang) if row["executed"] else t("no", lang)

    if not data["measurements"].empty:
        doc.add_heading(t("report_measurements_trend", lang), level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = t("measurement_date", lang)
        hdr[1].text = t("measurement_test", lang)
        hdr[2].text = t("measurement_score", lang)
        hdr[3].text = t("notes", lang)
        for _, row in data["measurements"].iterrows():
            r = table.add_row().cells
            r[0].text = str(row["measurement_date"])
            r[1].text = str(row["test_code"])
            r[2].text = f"{row['score']}/{row['max_score'] or '-'}"
            r[3].text = str(row.get("notes") or "")

    out = _filename(elder, "docx")
    doc.save(out)
    return str(out)


# ---------- PDF ----------
def elder_summary_pdf(conn, elder_id: int, days_back: int, lang: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    data = _summary_frame(conn, elder_id, days_back)
    elder = data["elder"]
    name_field = "name_he" if lang == "he" else "name_en"

    out = _filename(elder, "pdf")
    doc = SimpleDocTemplate(str(out), pagesize=A4,
                            rightMargin=20, leftMargin=20, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Title"], alignment=2 if lang == "he" else 0)
    body = ParagraphStyle("body", parent=styles["Normal"], alignment=2 if lang == "he" else 0, fontSize=10)

    story = []
    # Note: full RTL/Hebrew shaping in reportlab needs a Hebrew font; we leave
    # the system default which may show Hebrew but not perfectly shaped.
    story.append(Paragraph(f"{t('report_summary', lang)} - {elder['full_name']}", title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        f"{t('birth_date', lang)}: {elder.get('birth_date') or '-'} | "
        f"{t('room_number', lang)}: {elder.get('room_number') or '-'} | "
        f"{t('report_period_days', lang)}: {days_back}", body))
    story.append(Spacer(1, 12))

    story.append(Paragraph(
        f"<b>{t('items_scheduled', lang)}:</b> {data['total_items']} | "
        f"<b>{t('executed_total', lang)}:</b> {data['done_items']} | "
        f"<b>{t('completion_rate', lang)}:</b> {data['completion_rate']:.1f}%",
        body))
    story.append(Spacer(1, 16))

    if not data["items"].empty:
        rows = [[
            t("plan_for_date", lang).strip(),
            t("time_morning", lang).split()[0],
            t("nav_profile", lang),
            t("duration_min", lang),
            t("executed", lang),
        ]]
        for _, r in data["items"].iterrows():
            rows.append([
                str(r["plan_date"]),
                r["time_slot"] or "",
                str(r[name_field] or r["activity_code"])[:30],
                str(r["duration_min"] or ""),
                "✓" if r["executed"] else "—",
            ])
        tbl = Table(rows, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#444")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(tbl)

    doc.build(story)
    return str(out)


# ---------- EXCEL ----------
def elder_summary_xlsx(conn, elder_id: int, days_back: int, lang: str) -> str:
    data = _summary_frame(conn, elder_id, days_back)
    elder = data["elder"]
    out = _filename(elder, "xlsx")

    with pd.ExcelWriter(out, engine="openpyxl") as xw:
        pd.DataFrame([{
            t("elder_name", lang): elder["full_name"],
            t("birth_date", lang): elder.get("birth_date"),
            t("room_number", lang): elder.get("room_number"),
            t("report_period_days", lang): days_back,
            t("items_scheduled", lang): data["total_items"],
            t("executed_total", lang): data["done_items"],
            t("completion_rate", lang): f"{data['completion_rate']:.1f}%",
        }]).to_excel(xw, sheet_name="Summary", index=False)

        if not data["items"].empty:
            data["items"].to_excel(xw, sheet_name="Plan Items", index=False)
        if not data["measurements"].empty:
            data["measurements"].to_excel(xw, sheet_name="Measurements", index=False)
    return str(out)
