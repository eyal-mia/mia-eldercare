"""
ElderCare - Institution rules bank stored in a Word (.docx) file.

This is the single source of truth for the institution's free-text rules.
Each rule is one paragraph (one line). The optimizer reads this file on
EVERY run and applies whatever time-based rules it can parse; the UI reads
and writes it (add / delete / bulk-edit). The Word file lives next to the
knowledge banks so a manager can also open and edit it directly in Word.

If the file is missing it is created with a title and a set of default
rules so the system always has something to read.
"""

from __future__ import annotations
from pathlib import Path
import re

_NUM_PREFIX = re.compile(r"^\s*\d+\.\s*")

RULES_DOC_PATH = (
    Path(__file__).resolve().parent.parent / "knowledge_banks" / "rules_bank.docx"
)

DOC_TITLE = "בנק כללי המוסד - ElderCare"
DOC_SUBTITLE = (
    "כל שורה היא כלל אחד. כללי זמן (אין פעילות אחרי / לפני HH:MM) נאכפים "
    "אוטומטית על ידי מנוע התכנון. שאר הכללים מוצגים לצוות."
)

DEFAULT_RULES = [
    "אין פעילות אחרי 20:00",
    "אין פעילות לפני 08:00",
    "ארוחת ערב ב-18:00 בחדר האוכל המרכזי",
    "ביקור משפחה - עד 3 איש בו זמנית, בחדר הדייר/ת",
    "אין פעילות חיצונית ביום שבת",
    "חזרה לחדר עד 21:30",
    "אנשי צוות זמינים 24/7 במוקד 105",
]

# A marker so we can tell title/subtitle paragraphs apart from real rules.
_HEADER_TEXTS = {DOC_TITLE, DOC_SUBTITLE, ""}


def _ensure_exists() -> None:
    if RULES_DOC_PATH.exists():
        return
    write_rules(DEFAULT_RULES)


def read_rules() -> list[str]:
    """Return the list of rule lines from the Word file. Creates the file
    with defaults if it is missing or unreadable."""
    try:
        from docx import Document
    except Exception:
        # python-docx not available — fall back to defaults in memory
        return list(DEFAULT_RULES)

    if not RULES_DOC_PATH.exists():
        write_rules(DEFAULT_RULES)
        return list(DEFAULT_RULES)

    try:
        doc = Document(str(RULES_DOC_PATH))
    except Exception:
        # corrupt file — rebuild it
        write_rules(DEFAULT_RULES)
        return list(DEFAULT_RULES)

    rules: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text or text in _HEADER_TEXTS:
            continue
        # strip the leading "N. " numbering that write_rules adds
        text = _NUM_PREFIX.sub("", text).strip()
        if text:
            rules.append(text)
    return rules


def write_rules(rules: list[str]) -> None:
    """Rebuild the Word file from scratch with a title + one paragraph per rule."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return  # nothing we can do without python-docx

    RULES_DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading(DOC_TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sub = doc.add_paragraph(DOC_SUBTITLE)
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in sub.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("")  # spacer

    for i, rule in enumerate(rules, start=1):
        rule = (rule or "").strip()
        if not rule:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"{i}. {rule}")
        run.font.size = Pt(12)

    doc.save(str(RULES_DOC_PATH))


def add_rule(text: str) -> list[str]:
    """Append one rule and persist. Returns the new rules list."""
    text = (text or "").strip()
    rules = read_rules()
    if text and text not in rules:
        rules.append(text)
        write_rules(rules)
    return rules


def delete_rule(index: int) -> list[str]:
    """Delete the rule at position `index` (0-based) and persist."""
    rules = read_rules()
    if 0 <= index < len(rules):
        rules.pop(index)
        write_rules(rules)
    return rules


def set_rules_from_text(blob: str) -> list[str]:
    """Replace all rules from a newline-separated blob."""
    rules = [ln.strip() for ln in (blob or "").splitlines() if ln.strip()]
    write_rules(rules)
    return rules


def rules_as_text() -> str:
    """All rules joined by newlines — what the optimizer parser consumes."""
    return "\n".join(read_rules())


# read_rules() at import is not desirable; callers trigger creation lazily.
